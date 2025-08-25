from docx import Document
import sys

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract_docx.py <docx_file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    text = extract_text_from_docx(file_path)
    print(text)